from rest_framework import generics, status
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.decorators import api_view, permission_classes
from django.shortcuts import render
import json
from rest_framework.decorators import api_view, permission_classes
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from django.core.mail import send_mail
from django.conf import settings

from .models import SubscriptionProfile
from .serializers import SubscriptionProfileSerializer

from config.settings import YANDEX_DISK_TOKEN



from docx import Document
from io import BytesIO
from django.core.mail import EmailMessage
import os

from datetime import date
from .models import Account, SubscriptionProfile
from .serializers import AccountSerializer

def handle_post_request(request):
    try:
        # Retrieve the subscription profile for the current user
        subscription_profile = SubscriptionProfile.objects.get(user=request.user)
    except SubscriptionProfile.DoesNotExist:
        return {"detail": "Subscription profile not found."}
    
    # Handle POST request: Create or update today's account
    if request.method == 'POST':
        account, created = Account.objects.get_or_create(
            profile=subscription_profile,
            created=date.today(),
            defaults={'updated': 0}
        )
        if not created:
            # Update the existing account
            account.updated += 1
            account.save()
        
        # Serialize and return the account data
        serializer = AccountSerializer(account)
        return serializer.data

    # If the method is not POST, return a method not allowed response
    return {"detail": "Method not allowed."}







def generate_and_send_invoice(context, template_path, recipient_email):

    admins = ['storagehills@yandex.ru']

    message_body = f"""Спасибо, что пользуетесь нашим сервисом. 
    Выбранный тариф: {context['tariff']}
    Период подписки: {context['period']}
 
    Пожалуйста, оплатите подписку в ближайшее время. 
    С уважением, 
    Команда StorageHills."""

    message_subject = f"""Оплата StorageHills {context['inn']}. Счет {context['account_number']}"""

    # Load the template document
    doc = Document(template_path)

    # Replace placeholders with actual data in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in context.items():
            if f'{{{key}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{key}}}', str(value))

    # Replace placeholders with actual data in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in context.items():
                    if f'{{{key}}}' in cell.text:
                        cell.text = cell.text.replace(f'{{{key}}}', str(value),)
    # Save the modified document to a BytesIO object (in-memory file)
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)  # Move the cursor to the beginning of the file

    # Prepare the email
    email = EmailMessage(
        subject=message_subject,
        body=message_body,
        # Replace with your email or settings.DEFAULT_FROM_EMAIL
        to=[recipient_email , admins[0]],
    )
    
    
    # Attach the .docx file
    email.attach(f"Счет-{context['account_number']}.docx", doc_io.read(), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    
    # Send the email
    email.send()
    

    return True  # Optionally return True when the function completes successfully





@api_view(['POST'])  # Assuming POST request for sending emails
@permission_classes([IsAuthenticated])  # Require authentication
def payment_email(request):
    result = handle_post_request(request)
    context = request.data
    context['account_number'] = result['value']
    print(context)
    # Process the result further if needed
    print(result['value'])
    template_path = os.path.join(settings.MEDIA_ROOT, 'instance.docx')
    send_is = generate_and_send_invoice(context,template_path,request.user.email)
    if send_is : 
        return Response({"message": "письмо отправлено!"})
    return Response({"message": "письмо не отправлено"}, status=status.HTTP_400_BAD_REQUEST)

def react_app(request):
    return render(request, 'index.html')

class SubscriptionProfileDetailView(generics.RetrieveAPIView):
    serializer_class = SubscriptionProfileSerializer
    permission_classes = [IsAuthenticated]
    def get_object(self):
        # Получаем профиль подписки текущего аутентифицированного пользователя
        return self.request.user.subscription_profile
